import os
import csv
from docx import Document

def extract_text_from_docx(docx_path):
    """Extract all text from a Word document in correct order."""
    doc = Document(docx_path)
    text_content = []
    
    # Extract all paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_content.append(paragraph.text)
    
    # Extract all table content
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        text_content.append(paragraph.text)
    
    return '\n'.join(text_content)

def process_word_docs_to_csv(word_docs_folder, output_csv):
    """Process all Word documents in folder and save to CSV with two columns."""
    
    # Get all .docx files
    docx_files = [f for f in os.listdir(word_docs_folder) if f.endswith('.docx')]
    docx_files.sort()  # Sort alphabetically
    
    # Extract text from each document
    all_data = []
    for docx_file in docx_files:
        docx_path = os.path.join(word_docs_folder, docx_file)
        print(f"Processing: {docx_file}")
        text = extract_text_from_docx(docx_path)
        print(f"  Extracted {len(text)} characters")
        
        # Remove numbering from title (e.g., "1.4 Upper Limb" -> "Upper Limb")
        title = docx_file.replace('.docx', '')
        # Split by first space and number pattern to remove numbering
        import re
        title = re.sub(r'^\d+(\.\d+)*\s*', '', title)
        
        all_data.append((title, text))
    
    # Write to CSV with two columns
    with open(output_csv, 'w', newline='', encoding='utf-8') as csvfile:
        writer = csv.writer(csvfile)
        writer.writerow(['title', 'content'])  # Header
        
        for title, text in all_data:
            writer.writerow([title, text])
    
    print(f"Successfully processed {len(all_data)} documents to {output_csv}")

if __name__ == "__main__":
    word_docs_folder = "/Users/AdrianL/Documents/Medtrainer/Word docs"
    
    # Process y1s1
    print("Processing y1s1...")
    output_csv = "/Users/AdrianL/Documents/Medtrainer/lecture_notes_y1s1.csv"
    process_word_docs_to_csv(word_docs_folder, output_csv)
    
    # Process y1s2
    print("\nProcessing y1s2...")
    output_csv = "/Users/AdrianL/Documents/Medtrainer/lecture_notes_y1s2_new.csv"
    process_word_docs_to_csv(word_docs_folder, output_csv)
